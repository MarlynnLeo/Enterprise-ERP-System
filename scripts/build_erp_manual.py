from pathlib import Path
import argparse
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]

MANUALS = {
    "general": {
        "source": ROOT / "docs" / "ERP_操作手册与业务流程.md",
        "output": ROOT / "docs" / "KACON_ERP_MES_操作手册与业务流程_V1.0.docx",
        "header": "KACON ERP/MES | 操作手册",
        "title": "KACON ERP/MES\n操作手册与业务流程",
        "subtitle": "从主数据、业务单据到库存、生产、质量、财务和结账的统一操作指南",
        "meta": [
            ("版本", "V1.0"),
            ("整理日期", "2026-07-20"),
            ("适用范围", "桌面端、移动端、系统管理员、各业务岗位及运维人员"),
            ("编制依据", "当前代码、路由、权限、数据库迁移与自动化测试"),
        ],
        "callout": "本手册按系统当前实现整理。正式上线前，请由业务负责人用真实组织、仓库、会计科目和审批矩阵完成一次用户验收。",
        "contents_callout": "快速定位：业务人员从第 5 节开始；主管重点阅读第 6、9 节；管理员和运维人员重点阅读第 4、7、11、12 节。",
        "subject": "ERP/MES 系统操作、端到端业务流程、上线初始化和运维",
        "comments": "面向系统管理员、业务岗位和运维人员的综合操作说明",
        "page_break_sections": {"附录 B：上线验收核心场景"},
    },
    "department": {
        "source": ROOT / "docs" / "ERP_部门业务操作手册.md",
        "output": ROOT / "docs" / "KACON_ERP_MES_部门业务操作手册_V1.0.docx",
        "header": "KACON ERP/MES | 部门业务手册",
        "title": "KACON ERP/MES\n部门业务操作手册",
        "subtitle": "给各部门同事的日常录入、审批、执行、交接与异常处理指南",
        "meta": [
            ("版本", "V1.0"),
            ("整理日期", "2026-07-20"),
            ("适用人员", "销售、采购、仓库、计划、生产、质量、财务、设备、人事及主管"),
            ("使用目的", "统一部门操作、单据填写、审批交接和业务闭环"),
        ],
        "callout": "本手册只说明业务操作。各部门应结合公司授权、合同、质量和财务制度执行，系统中没有的特殊业务请先向主管确认。",
        "contents_callout": "快速定位：所有人先看第 3、4 节；销售看第 6 节；采购看第 7 节；仓库看第 8 节；计划/生产看第 9 节；质量看第 10 节；财务看第 11 节。",
        "subject": "ERP/MES 部门业务操作、审批、交接和异常处理",
        "comments": "面向部门同事的业务操作说明",
        "bottom_margin": 0.6,
        "compact_numbered_sections": {
            "附录 B：业务单据交接模板",
            "附录 C：业务培训演练场景",
        },
    },
    "technical": {
        "source": ROOT / "docs" / "ERP_技术开发与交接文档.md",
        "output": ROOT / "docs" / "KACON_ERP_MES_技术开发与交接文档_V1.0.docx",
        "header": "KACON ERP/MES | 技术交接文档",
        "title": "KACON ERP/MES\n技术开发与交接文档",
        "subtitle": "面向后端、桌面端、移动端、数据库、部署和运维接手人员",
        "meta": [
            ("版本", "V1.0"),
            ("整理日期", "2026-07-20"),
            ("适用对象", "后端、桌面端、移动端、数据库、部署和运维接手人员"),
            ("文档定位", "以当前代码和部署文件为准的开发、发布、排障与回滚指南"),
        ],
        "callout": "这份文档描述当前仓库真实实现。接手开发前先阅读已知边界与高风险规则，不要根据页面名称推断后端一定存在对应工作流或接口。",
        "contents_callout": "快速定位：先看第 2、3、4 节；本地开发看第 5、6 节；测试看第 7 节；运维发布看第 8、9、10 节。",
        "subject": "ERP/MES 技术架构、开发约定、部署运维和接手验收",
        "comments": "面向后端、前端、数据库、部署和运维接手人员的技术交接说明",
    },
}

BLUE = "2E74B5"
DEEP_BLUE = "1F4D78"
MUTED = "5B6573"
LIGHT_BLUE = "E8EEF5"
LIGHT_CALLOUT = "F3F7FB"
WHITE = "FFFFFF"
BLACK = "1F2933"
TABLE_WIDTH = 9360


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    total = sum(widths)
    if total != TABLE_WIDTH:
        widths[-1] += TABLE_WIDTH - total
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row_idx, row in enumerate(table.rows):
        if row_idx == 0:
            tr_pr = row._tr.get_or_add_trPr()
            header = OxmlElement("w:tblHeader")
            header.set(qn("w:val"), "true")
            tr_pr.append(header)
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.1


def set_run_font(run, name="Calibri", size=11, color=BLACK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_border(paragraph, color=BLUE, size="12", space="8"):
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    set_run_font(run, size=9, color=MUTED)


def add_inline_markdown(paragraph, text):
    # Keep the parser intentionally small and deterministic for this maintained manual.
    pattern = re.compile(r"(\*\*.*?\*\*|`.*?`|\[.*?\]\(.*?\))")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            set_run_font(run)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=10, color=DEEP_BLUE)
        else:
            label, _, url = token[1:-1].partition("](")
            run = paragraph.add_run(label)
            set_run_font(run, color=BLUE, bold=True)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run)


def add_body_paragraph(doc, text, style="Normal"):
    p = doc.add_paragraph(style=style)
    add_inline_markdown(p, text)
    return p


def create_numbering(doc):
    """Create a fresh single-level decimal list so each Markdown list restarts at 1."""
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract.append(multi_level)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1.")
    level.append(level_text)
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "space")
    level.append(suffix)
    p_pr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_numbered_paragraph(doc, text, num_id):
    p = add_body_paragraph(doc, text, style="List Number")
    p_pr = p._p.get_or_add_pPr()
    existing = p_pr.find(qn("w:numPr"))
    if existing is not None:
        p_pr.remove(existing)
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_pr.append(ilvl)
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(num_id_node)
    p_pr.append(num_pr)
    return p


def add_callout(doc, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_WIDTH])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_CALLOUT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    add_inline_markdown(p, text)
    return table


def parse_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        rows.append(cells)
        i += 1
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in rows[1]):
        rows.pop(1)
    return rows, i


def add_table(doc, rows):
    if not rows:
        return
    cols = max(len(row) for row in rows)
    rows = [row + [""] * (cols - len(row)) for row in rows]
    table = doc.add_table(rows=len(rows), cols=cols)
    if cols == 1:
        widths = [TABLE_WIDTH]
    elif cols == 2:
        widths = [2500, TABLE_WIDTH - 2500]
    elif cols == 3:
        widths = [1800, 3400, TABLE_WIDTH - 5200]
    else:
        widths = [TABLE_WIDTH // cols] * cols
        widths[-1] += TABLE_WIDTH - sum(widths)
    set_table_geometry(table, widths)
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            if r_idx == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in list(p.runs):
                run._element.getparent().remove(run._element)
            add_inline_markdown(p, value)
            if r_idx == 0:
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(DEEP_BLUE)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DEEP_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name, left, hanging in [("List Bullet", 0.375, 0.188), ("List Number", 0.375, 0.188)]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(left)
        style.paragraph_format.first_line_indent = Inches(-hanging)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_cover(doc, config):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run(config["header"])
    set_run_font(run, size=9, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("内部使用 | V1.0 | 第 ")
    set_run_font(run, size=9, color=MUTED)
    add_page_field(footer)
    run = footer.add_run(" 页")
    set_run_font(run, size=9, color=MUTED)

    for _ in range(5):
        doc.add_paragraph().paragraph_format.space_after = Pt(12)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("ERP / MES OPERATOR GUIDE")
    set_run_font(run, size=11, color=BLUE, bold=True)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(14)
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run(config["title"])
    set_run_font(run, size=27, color=DEEP_BLUE, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    run = subtitle.add_run(config["subtitle"])
    set_run_font(run, size=12, color=MUTED)
    set_paragraph_border(subtitle, color=BLUE, size="8", space="10")

    meta = doc.add_table(rows=4, cols=2)
    set_table_geometry(meta, [2400, TABLE_WIDTH - 2400])
    entries = config["meta"]
    for i, (label, value) in enumerate(entries):
        set_cell_shading(meta.cell(i, 0), LIGHT_BLUE)
        p = meta.cell(i, 0).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(label)
        set_run_font(r, size=10.5, color=DEEP_BLUE, bold=True)
        p2 = meta.cell(i, 1).paragraphs[0]
        r2 = p2.add_run(value)
        set_run_font(r2, size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    add_callout(doc, config["callout"])
    doc.add_page_break()


def add_contents(doc, source_lines, config):
    p = doc.add_paragraph(style="Heading 1")
    p.add_run("目录")
    num_id = create_numbering(doc)
    for line in source_lines:
        if line.startswith("## "):
            text = re.sub(r"^\d+\.\s*", "", line[3:].strip())
            add_numbered_paragraph(doc, text, num_id)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    add_callout(doc, config["contents_callout"])
    doc.add_page_break()


def build(config):
    lines = config["source"].read_text(encoding="utf-8").splitlines()
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(config.get("bottom_margin", 0.75))
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    configure_styles(doc)
    add_cover(doc, config)
    add_contents(doc, lines, config)

    # The cover already contains the Markdown title and metadata block.
    i = next((idx for idx, line in enumerate(lines) if line.startswith("## ")), 0)
    current_section = ""
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        if line.startswith("# "):
            i += 1
            continue
        if line.startswith("## "):
            current_section = line[3:].strip()
            if current_section in config.get("page_break_sections", set()):
                doc.add_page_break()
            p = doc.add_paragraph(style="Heading 1")
            p.add_run(current_section)
            i += 1
            continue
        if line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            p.add_run(line[4:].strip())
            i += 1
            continue
        if line.startswith("#### "):
            p = doc.add_paragraph(style="Heading 3")
            p.add_run(line[5:].strip())
            i += 1
            continue
        if line.startswith("> "):
            add_callout(doc, line[2:].strip())
            i += 1
            continue
        if line.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue
        if re.match(r"^- \[[ xX]\] ", line):
            text = re.sub(r"^- \[[ xX]\] ", "[ ] ", line)
            add_body_paragraph(doc, text, style="List Bullet")
            i += 1
            continue
        if line.startswith("- "):
            add_body_paragraph(doc, line[2:].strip(), style="List Bullet")
            i += 1
            continue
        if re.match(r"^\d+\. ", line):
            num_id = create_numbering(doc)
            compact_numbered = current_section in config.get("compact_numbered_sections", set())
            while i < len(lines) and re.match(r"^\d+\. ", lines[i].rstrip()):
                text = re.sub(r"^\d+\. ", "", lines[i].rstrip())
                p = add_numbered_paragraph(doc, text, num_id)
                if compact_numbered:
                    p.paragraph_format.space_after = Pt(1)
                    p.paragraph_format.line_spacing = 1.2
                i += 1
            continue
        if line == "---":
            p = doc.add_paragraph()
            set_paragraph_border(p, color="D7E3F0", size="4", space="1")
            i += 1
            continue
        add_body_paragraph(doc, line)
        i += 1

    # Ensure the document has a useful core property set.
    doc.core_properties.title = config["title"].replace("\n", " ")
    doc.core_properties.subject = config["subject"]
    doc.core_properties.author = "KACON ERP"
    doc.core_properties.comments = config["comments"]
    config["output"].parent.mkdir(parents=True, exist_ok=True)
    doc.save(config["output"])
    print(config["output"])


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Build the ERP manuals from Markdown.")
    parser.add_argument(
        "manual",
        nargs="?",
        choices=["general", "department", "all"],
        default="all",
    )
    args = parser.parse_args()
    names = MANUALS if args.manual == "all" else [args.manual]
    for name in names:
        build(MANUALS[name])
